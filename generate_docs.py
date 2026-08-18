from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Text Correction & Spell Checking Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "This project is a Full-Stack Machine Learning application designed to automatically rectify "
        "spelling and grammatical errors in textual data. It implements two distinct approaches: a "
        "Classical NLP approach (using SymSpell and TextBlob) and a Modern Deep Learning approach "
        "(fine-tuned T5 Transformer model). Both models are evaluated, and the best performing model is "
        "served via a FastAPI backend and consumed by a modern React-based UI."
    )
    
    # Architecture
    doc.add_heading('2. Architecture', level=1)
    doc.add_paragraph(
        "• Frontend: Built with React and Vite. It provides a premium, responsive UI where users can type "
        "text and receive real-time corrections.\n"
        "• Backend: Developed using FastAPI. It exposes a REST API endpoint (/api/correct) and dynamically "
        "loads the optimal model based on an evaluation phase.\n"
        "• Evaluation Module: A script that runs a predefined test dataset through both engines and calculates "
        "the Word Error Rate (WER). The model with the lowest WER is selected as the active model."
    )
    
    # Classical Method
    doc.add_heading('3. Classical NLP Method', level=1)
    doc.add_paragraph(
        "The classical approach relies on SymSpell and TextBlob.\n\n"
        "SymSpell: Uses a Symmetric Delete spelling correction algorithm. It pre-calculates all possible "
        "delete combinations within a given edit distance, allowing for extremely fast dictionary lookups "
        "in O(1) time.\n\n"
        "TextBlob: A simpler library used as a fallback for grammatical heuristics."
    )
    
    # Modern Method
    doc.add_heading('4. Modern Fine-Tuned Method', level=1)
    doc.add_paragraph(
        "The modern approach utilizes the Hugging Face Transformers library, specifically a T5 "
        "(Text-to-Text Transfer Transformer) model fine-tuned for grammar correction (e.g., "
        "vennify/t5-base-grammar-correction). T5 treats all NLP tasks as a text-to-text problem, making "
        "it highly effective for generating corrected versions of grammatically incorrect input."
    )
    
    # Evaluation Metrics
    doc.add_heading('5. Evaluation & Metrics', level=1)
    doc.add_paragraph(
        "The models were evaluated using the Word Error Rate (WER) metric via the jiwer library. "
        "WER calculates the minimum number of substitutions, deletions, and insertions required to "
        "change the predicted text into the reference (correct) text, normalized by the total number of "
        "words in the reference text. A lower WER indicates higher accuracy."
    )
    
    # Interview Talking Points
    doc.add_heading('6. Interview Talking Points', level=1)
    doc.add_paragraph(
        "• Why SymSpell? Mention its O(1) lookup speed compared to Peter Norvig's approach which "
        "generates candidate edits at runtime.\n"
        "• Why T5? Emphasize the text-to-text framework which is ideal for grammar correction where "
        "both input and output are sequences of text.\n"
        "• Why evaluate? Stress that in MLOps, models should be empirically tested against a dataset "
        "before deployment."
    )
    
    doc.save('Text_Correction_Interview_Prep.docx')
    print("Interview prep document generated successfully!")

if __name__ == '__main__':
    create_document()
